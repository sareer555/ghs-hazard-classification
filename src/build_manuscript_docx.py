"""
CONVERT THE MANUSCRIPT TO A WORD DOCUMENT
=========================================
Elsevier asks for an editable source file - "a PDF is not an acceptable source
file" - and the manuscript is generated as plain text. Doing the conversion by
hand means re-doing it by hand every time a number changes, and hand conversion
is where section numbering and hanging indents get quietly lost.

The generated text is hard-wrapped to 79 columns. Pasting it into Word directly
would preserve those line breaks as real breaks, so the first job here is to
undo the wrapping and recover genuine paragraphs; the second is to map the
plain-text structure onto Word heading styles so the section numbering survives
and a table of contents is possible.

Output: publication_materials/manuscript/GHS_manuscript.docx

Author : Sareer Ahmad
"""

import os
import re
import sys

import docx
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from ghs_config import DIR_PUB

MANUSCRIPT_DIR = os.path.join(DIR_PUB, "manuscript")
SOURCE = os.path.join(MANUSCRIPT_DIR, "FULL_MANUSCRIPT.txt")
OUTPUT = os.path.join(MANUSCRIPT_DIR, "GHS_manuscript.docx")

# Labels on the title page. They are headings in the plain text but should read
# as field labels in Word, not as document sections.
TITLE_PAGE_LABELS = {"TITLE", "RUNNING TITLE", "AUTHOR", "AFFILIATION",
                     "KEYWORDS", "DISCLAIMER", "BIOGRAPHICAL NOTE"}

SECTION_NUMBER = re.compile(r"^(\d+)\.\s+(.+)$")
SUBSECTION_NUMBER = re.compile(r"^(\d+\.\d+)\.\s+(.+)$")
REFERENCE_ENTRY = re.compile(r"^\((\d+)\)\s")
ALL_CAPS_HEADING = re.compile(r"^[A-Z][A-Z0-9 ,()\-/&]{3,}$")


def is_rule(line):
    """True for the rows of dashes or equals signs used as underlines."""
    stripped = line.strip()
    return bool(stripped) and set(stripped) <= set("=-")


COLUMNAR = re.compile(r"\S {3,}\S")


def looks_like_table_header(line):
    """
    True for a column header row rather than a section heading.

    Both are underlined with dashes in the plain text, so the underline alone
    cannot tell them apart. A column header has text separated into columns by
    runs of two or more spaces; a heading is ordinary prose.
    """
    return bool(COLUMNAR.search(line.strip()))


def classify(line, previous_blank, next_line, after_conclusions):
    """
    Decide what a single line of the plain-text manuscript is.

    Headings are recognised by the rule of dashes or equals signs beneath them,
    which is how the generated text marks them, rather than by capitalisation.
    Case-based detection missed the two headings that are not upper case - the
    references heading and the generative-AI declaration, the latter being one
    Elsevier requires to be a section of its own.

    Returns one of: rule, blank, subsection, section, label, reference,
    indented, body.
    """
    stripped = line.strip()
    if not stripped:
        return "blank"
    if is_rule(line):
        return "rule"
    if stripped in TITLE_PAGE_LABELS:
        return "label"

    # A numbered top-level heading stands on its own, underlined or not.
    if SECTION_NUMBER.match(stripped) and not SUBSECTION_NUMBER.match(stripped):
        return "section"
    # The abstract heading is neither numbered nor underlined, so it needs
    # naming. Checked against the parsed output: it is the only one.
    if stripped == "ABSTRACT":
        return "section"

    underlined = next_line is not None and is_rule(next_line)
    if underlined and not looks_like_table_header(line):
        if stripped.startswith("REFERENCES"):
            return "section"
        # Back matter sits between the conclusions and the references, and its
        # sections are top level even though they carry a dash underline.
        return "section" if after_conclusions else "subsection"

    # Methods subsections are numbered and carry no underline.
    if SUBSECTION_NUMBER.match(stripped):
        return "subsection"
    if REFERENCE_ENTRY.match(stripped):
        return "reference"
    if line.startswith("    ") and previous_blank:
        return "indented"
    return "body"


def split_columns(line, expected):
    """
    Split one row of a fixed-width ASCII table into its fields.

    Fields are separated by runs of two or more spaces. The data rows put the
    pictogram code and the hazard name in one field separated by a single
    space, so a row one field short of the header is split again on its first
    space. Anything else is returned as it is, padded, rather than guessed at.
    """
    fields = re.split(r" {2,}", line.strip())
    if len(fields) == expected - 1 and " " in fields[0]:
        head, _, rest = fields[0].partition(" ")
        fields = [head, rest] + fields[1:]
    if len(fields) < expected:
        fields += [""] * (expected - len(fields))
    return fields[:expected]


def read_table(lines, start):
    """
    Read the ASCII table whose header is at `start`.

    Returns (header_fields, data_rows, index_after_the_table), or None when
    what follows the rule is not a table after all. The plain text marks a
    table as a columnar header, a rule of dashes, then rows until a blank line.
    """
    header = re.split(r" {2,}", lines[start].strip())
    if len(header) < 2:
        return None
    rows, position = [], start + 2          # skip the header and its rule
    while position < len(lines) and lines[position].strip():
        if is_rule(lines[position]):
            break
        rows.append(split_columns(lines[position], len(header)))
        position += 1
    return (header, rows, position) if rows else None


def parse(text):
    """
    Turn the wrapped plain text into a list of (kind, content) blocks.

    Consecutive body lines are joined back into one paragraph, which is the
    whole point: the source is wrapped for reading in a terminal, and those
    line breaks must not survive into Word. The ASCII tables are recovered as
    structured rows: left as text they reflow into a run-on line, which is
    exactly the defect that made the results unreadable.
    """
    lines = text.split("\n")
    blocks = []
    buffer, buffer_kind = [], None

    def flush():
        """Emit whatever paragraph has been accumulating."""
        if buffer:
            blocks.append((buffer_kind, " ".join(" ".join(buffer).split())))
            buffer.clear()

    previous_blank = True
    after_conclusions = False
    skip_until = 0
    for position, line in enumerate(lines):
        if position < skip_until:
            continue
        next_line = lines[position + 1] if position + 1 < len(lines) else None

        # A columnar line above a rule of dashes opens an ASCII table.
        if (line.strip() and next_line is not None and is_rule(next_line)
                and looks_like_table_header(line)):
            table = read_table(lines, position)
            if table:
                header, rows, skip_until = table
                flush()
                blocks.append(("table", (header, rows)))
                buffer_kind = None
                previous_blank = False
                continue
        kind = classify(line, previous_blank, next_line, after_conclusions)
        previous_blank = kind == "blank"

        # Everything after the conclusions and before the references is back
        # matter, whose sections are top level.
        if kind == "section":
            heading = line.strip()
            if heading.startswith("5. CONCLUSIONS"):
                after_conclusions = True
            elif heading.startswith("REFERENCES"):
                after_conclusions = False

        if kind in ("rule",):
            continue
        if kind == "blank":
            flush()
            buffer_kind = None
            continue
        if kind in ("section", "subsection", "label"):
            flush()
            blocks.append((kind, line.strip()))
            buffer_kind = None
            continue
        if kind == "reference":
            flush()
            buffer_kind = "reference"
            buffer.append(line.strip())
            continue
        # body or indented: keep accumulating, but a change of kind starts a
        # new paragraph so an indented block is not merged into running prose.
        if buffer_kind in (None, kind) or (buffer_kind == "reference"):
            buffer_kind = buffer_kind or kind
            buffer.append(line.strip())
        else:
            flush()
            buffer_kind = kind
            buffer.append(line.strip())
    flush()
    return blocks


def add_line_numbers(section):
    """
    Turn on continuous line numbering.

    Reviewers refer to a manuscript by line number, and python-docx has no API
    for this, so the section properties are edited directly.
    """
    properties = section._sectPr
    numbering = OxmlElement("w:lnNumType")
    numbering.set(qn("w:countBy"), "1")
    numbering.set(qn("w:restart"), "continuous")
    numbering.set(qn("w:distance"), "360")      # twentieths of a point
    properties.append(numbering)


def add_page_numbers(section):
    """Put 'Page N' in the footer, as a field Word evaluates itself."""
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for instruction, kind in (("begin", "w:fldChar"), ("PAGE", "w:instrText"),
                              ("end", "w:fldChar")):
        element = OxmlElement(kind)
        if kind == "w:fldChar":
            element.set(qn("w:fldCharType"), instruction)
        else:
            element.set(qn("xml:space"), "preserve")
            element.text = " PAGE "
        run._r.append(element)


def configure(document):
    """Set the page, the base font and the spacing the whole document uses."""
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)   # A4
    for margin in ("left_margin", "right_margin"):
        setattr(section, margin, Cm(2.5))
    section.top_margin = section.bottom_margin = Cm(2.5)
    add_line_numbers(section)
    add_page_numbers(section)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # East Asian font must be set too or Word substitutes for some characters.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_after = Pt(0)

    for name, size in (("Heading 1", 14), ("Heading 2", 12.5),
                       ("Heading 3", 12)):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.keep_with_next = True


def build(blocks):
    """Write the Word document from the parsed blocks."""
    document = docx.Document()
    configure(document)

    TABLE_CAPTION = re.compile(r"^Table \d+\.\s")

    def add_table(header, rows):
        """Render one parsed ASCII table as a real Word table."""
        table = document.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for cell, name in zip(table.rows[0].cells, header):
            cell.text = ""
            run = cell.paragraphs[0].add_run(name)
            run.bold = True
            run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.line_spacing_rule =                 WD_LINE_SPACING.SINGLE
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        for values in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                cell.text = ""
                run = cell.paragraphs[0].add_run(value)
                run.font.size = Pt(9)
                cell.paragraphs[0].paragraph_format.line_spacing_rule =                     WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        # A table butted against the next paragraph is unreadable.
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for kind, content in blocks:
        if kind == "table":
            add_table(*content)
            continue
        # The text file's own banner, not a manuscript section; the title page
        # content follows it either way.
        if content == "TITLE PAGE":
            continue
        if kind == "section":
            document.add_paragraph(content, style="Heading 1")
        elif kind == "subsection":
            document.add_paragraph(content, style="Heading 2")
        elif kind == "label":
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.add_run(content.title())
        elif kind == "reference":
            paragraph = document.add_paragraph(content)
            # Hanging indent, so the number sits proud of the wrapped text.
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.first_line_indent = Cm(-1.0)
            paragraph.paragraph_format.line_spacing_rule = \
                WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(6)
        elif kind == "indented":
            paragraph = document.add_paragraph(content)
            paragraph.paragraph_format.left_indent = Cm(1.0)
        elif TABLE_CAPTION.match(content):
            paragraph = document.add_paragraph()
            number, _, rest = content.partition(". ")
            bold = paragraph.add_run(number + ". ")
            bold.bold = True
            paragraph.add_run(rest)
            paragraph.paragraph_format.line_spacing_rule =                 WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                run.font.size = Pt(10)
        else:
            document.add_paragraph(content)

    return document


def main():
    """Convert the manuscript and report what was produced."""
    print("=" * 78)
    print("CONVERTING THE MANUSCRIPT TO WORD")
    print("=" * 78)

    if not os.path.exists(SOURCE):
        raise SystemExit(f"Manuscript not found: {SOURCE}\n"
                         f"Run manuscript_sections.py first.")

    with open(SOURCE, encoding="utf-8") as fh:
        blocks = parse(fh.read())

    document = build(blocks)
    document.save(OUTPUT)

    counts = {}
    for kind, _ in blocks:
        counts[kind] = counts.get(kind, 0) + 1
    print("\n   blocks written")
    for kind in sorted(counts):
        print(f"      {kind:<12}{counts[kind]:>5}")

    # Table blocks carry (header, rows), not text; word count is prose only.
    words = sum(len(c.split()) for k, c in blocks if k != "table")
    print(f"\n   {words:,} words")
    print(f"   {OUTPUT}")
    print(f"   {os.path.getsize(OUTPUT) / 1024:,.0f} KB")
    print("\nA4, single column, double spaced, continuous line numbers,")
    print("Times New Roman 12 pt, page numbers in the footer.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
